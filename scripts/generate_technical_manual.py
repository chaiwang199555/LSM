#!/usr/bin/env python3
"""
Generate LSM Technical Manual (Word .docx)
Similar structure to CLM Technical Description: TOC, chapters per physics process,
introductions, governing equations, numerical methods, and parameter tables.
"""

from __future__ import annotations

import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

OUTPUT = Path(__file__).resolve().parent.parent / "LSM_TECHNICAL_MANUAL.docx"
VERSION = "2.1"
DATE = "2026-06"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_doc_styles(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(11)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15
    pf.space_after = Pt(6)

    for level, size in [(1, 16), (2, 14), (3, 12)]:
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Arial"
        hs.font.size = Pt(size)
        hs.font.bold = True
        hs.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)


def add_page_number_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    run._r.append(fld1)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    run._r.append(instr)
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld2)


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    run._r.append(instr)
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)
    note = doc.add_paragraph(
        "（打开文档后，在目录区域右键选择“更新域”以生成页码。）"
    )
    note.runs[0].italic = True
    note.runs[0].font.size = Pt(9)


def add_para(doc: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def add_equation(doc: Document, eq: str, number: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    if number:
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(6.0), WD_ALIGN_PARAGRAPH.RIGHT)
        run = p.add_run(eq)
        run.italic = True
        run.font.name = "Cambria Math"
        run.font.size = Pt(11)
        p.add_run("\t")
        num_run = p.add_run(f"({number})")
        num_run.font.size = Pt(11)
    else:
        run = p.add_run(eq)
        run.italic = True
        run.font.name = "Cambria Math"
        run.font.size = Pt(11)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    doc.add_paragraph()


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Content sections
# ---------------------------------------------------------------------------

def build_title_page(doc: Document) -> None:
    for _ in range(6):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("单点陆面模式\nLand Surface Model (LSM)\n技术手册")
    r.bold = True
    r.font.size = Pt(22)
    r.font.name = "Arial"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(
        f"Technical Description of Physics and Numerics\n"
        f"版本 {VERSION}  |  {DATE}"
    )
    sr.font.size = Pt(12)
    sr.italic = True

    doc.add_paragraph()
    inst = doc.add_paragraph()
    inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    inst.add_run(
        "面向站点尺度陆气相互作用、植被–土壤–碳氮磷耦合研究\n"
        "/home/chai/class/data/LSM"
    ).font.size = Pt(10)

    add_page_break(doc)


def build_chapter1(doc: Document) -> None:
    doc.add_heading("第1章  引言", level=1)

    doc.add_heading("1.1  文档目的", level=2)
    add_para(
        doc,
        "本技术手册详细描述单点陆面模式（LSM）的物理过程、控制方程、数值求解方案、"
        "状态变量与通量诊断量，以及各物理方案之间的耦合关系。文档结构参考 CLM "
        "（Community Land Model）技术说明书的组织方式：按物理过程分章阐述，每章包含"
        "过程简介、控制方程、离散化与实现要点、相关参数及参考文献。"
    )

    doc.add_heading("1.2  模式特征", level=2)
    add_table(
        doc,
        ["项目", "说明"],
        [
            ["空间维度", "单点（0-D），无水平异质性与平流输送"],
            ["时间积分", "显式前向差分，默认 Δt = 1800 s（30 min）"],
            ["语言与精度", "Fortran 2008，双精度 real64"],
            ["强迫数据", "文本或 NetCDF 格式逐时步大气边界条件（含 cos_sza、sw_beam_frac）"],
            ["物理框架", "Noah-MP 风格多物理开关（istomatal、icanopy、irad 等）"],
            ["碳循环", "可选（carbon_on），含 GPP、土壤呼吸、简化 CNP"],
            ["适用区域", "青藏高原及一般陆地站点快速原型研究"],
        ],
    )

    doc.add_heading("1.3  物理过程耦合概览", level=2)
    add_para(
        doc,
        "模式在每个时间步内按松耦合顺序执行：GCM 边界耦合 → 光合适应 → 地表能量平衡"
        "（Noah-MP 序贯冠层/地面求解）→ 植物水力 → 土壤温度 → 冻土相变 → 积雪"
        "→ [可选] 双叶光合/呼吸/碳库/CNP → 土壤水分 → 守恒检查。"
        "当 icanopy=2 时，能量平衡在冠层温度 Tc 与地面温度 Ts 上分别迭代，短波采用 "
        "Noah-MP 两流方案，长波采用 Noah-MP 层间交换系数。"
    )
    add_bullet(doc, "冠层辐射（mod_radtran）：Noah 两流短波 + Noah 长波层间交换")
    add_bullet(doc, "湍流交换：Monin-Obukhov 相似理论求感热 H 与空气动力学阻力 ra")
    add_bullet(doc, "地表能量平衡：序贯求解 Tc（冠层）与 Ts（地面），组装 H、LE、G")
    add_bullet(doc, "气孔与光合：Jarvis 或 Medlyn 气孔导度，Farquhar C3 光合")
    add_bullet(doc, "植物水力：PHS 链式水势与木质部脆弱性")
    add_bullet(doc, "土壤热传导：顶层显式能量收支，深层松弛扩散")
    add_bullet(doc, "土壤水分：Bucket 或 Richards + van Genuchten")
    add_bullet(doc, "土壤碳与养分：Q10/微生物呼吸，简化 CASA-CNP")


def build_chapter2(doc: Document) -> None:
    doc.add_heading("第2章  模式架构与数值框架", level=1)

    doc.add_heading("2.1  程序调用链", level=2)
    add_para(doc, "主程序 main.f90 读取 namelist 后调用 mod_driver::run_lsm 完成时间积分：")
    add_bullet(doc, "read_namelist() [mod_io]")
    add_bullet(doc, "load_forcing() → init_state() → init_output()")
    add_bullet(doc, "循环 nspinup + 1 次强迫序列（前 nspinup 次为 spin-up，不写输出）")
    add_bullet(doc, "每步：couple_in → update_acclimation → solve_energy_balance → update_plant_hydro "
               "→ update_soil_temperature → update_phase_change → update_snow → [carbon] "
               "→ update_soil_water → couple_out → check_balances")

    doc.add_heading("2.2  源码模块", level=2)
    add_table(
        doc,
        ["模块", "职责"],
        [
            ["mod_radtran", "Noah 两流短波、Noah 长波交换、PAR 阴阳叶分配"],
            ["mod_surface", "Noah-MP 序贯能量平衡（双叶 Tc/Ts）"],
            ["mod_canopy", "双叶气孔与冠层光合"],
            ["mod_ncio", "NetCDF 强迫/土壤初值/输出 I/O"],
            ["mod_snow / mod_permafrost", "积雪与冻土相变（MVP）"],
            ["mod_gcm_coupling", "GCM 边界占位耦合（irad=2）"],
        ],
    )

    doc.add_heading("2.3  核心状态变量", level=2)
    add_table(
        doc,
        ["类型", "主要字段", "单位/说明"],
        [
            ["t_forcing", "SW, LW, Ta, P, WS, PA, CO2, RH, cos_sza, sw_beam_frac", "RH 读入后派生 VPD"],
            ["t_state", "Ts, Tc, LAI, LAI_sun/shade, beta, W, psi_*, Tsoil, theta, snow_*", "水势 MPa"],
            ["t_param", "z0, albedo, clumping, gs_max, g0/g1, Vcmax, Ksat, ...", "站点参数"],
            ["t_flux", "Rn, H, LE, G, SW_abs_*, PAR_sun/shade, GPP_sun/shade, albedo_eff", "通量诊断"],
        ],
    )

    doc.add_heading("2.4  物理方案开关", level=2)
    add_table(
        doc,
        ["namelist 变量", "选项", "模块", "物理含义"],
        [
            ["istomatal", "1=Jarvis, 2=Medlyn", "mod_conductance", "气孔导度方案"],
            ["icanopy", "1=大叶, 2=双叶", "mod_surface, mod_canopy", "冠层结构与能量/光合"],
            ["irad", "0=离线, 1=两流, 2=GCM", "mod_radtran", "短波/长波辐射方案"],
            ["ihydro", "1=beta, 2=PHS", "mod_planthydro", "蒸腾水分胁迫"],
            ["isoilwater", "1=Bucket, 2=Richards", "mod_soil_water", "土壤水分运动"],
            ["isnow", "0=关, 1=开", "mod_snow", "积雪过程"],
            ["ifrost", "0=关, 1=开", "mod_permafrost", "冻土相变"],
            ["isoilcarbon", "1=Q10, 2=微生物", "mod_soil_carbon", "土壤呼吸与碳库"],
            ["icnp", "0/1/2/3", "mod_cnp", "碳氮磷养分耦合"],
            ["eeo_on", ".true./.false.", "mod_acclimation", "EEO 光合适应"],
        ],
    )

    doc.add_heading("2.5  Spin-up 策略", level=2)
    add_para(
        doc,
        "驱动程序将强迫时间序列循环复用 nspinup + 1 次。前 nspinup 次循环仅更新模式状态、"
        "不写磁盘输出；最后一次为生产运行。强迫索引通过取模实现循环，适合用短期（如 1 天）"
        "强迫数据做多日 spin-up，但尚未实现多年气候序列自动拼接。"
    )


def build_chapter3(doc: Document) -> None:
    doc.add_heading("第3章  辐射与湿度", level=1)
    doc.add_heading("3.1  过程简介", level=2)
    add_para(
        doc,
        "辐射计算分为两层：mod_radiation 提供湿度、比湿与离线净辐射；mod_radtran 在 "
        "irad≥1 时实现 Noah-MP 风格冠层辐射传输。短波采用 Dickinson/Niu-Yang 两流方案"
        "（可见光 48% + 近红外两波段），长波采用 Noah-MP 冠层–地面层间交换系数。"
        "强迫文件可提供 cos(SZA) 与直射比例 sw_beam_frac；缺省时由总短波推断。"
    )

    doc.add_heading("3.2  短波几何与两流传输（irad = 1）", level=2)
    add_para(doc, "直射/散射拆分：")
    add_equation(doc, "SW_dir = SW_total · sw_beam_frac", "3.1")
    add_equation(doc, "SW_dif = SW_total − SW_dir", "3.2")
    add_para(doc, "对每个波段（vis, nir）调用 noah_twostream_unit，得到冠层吸收、地表透射与有效反照率 albedo_eff。")
    add_bullet(doc, "cos_sza：太阳天顶角余弦，夜间取 0.001")
    add_bullet(doc, "sw_beam_frac：直射束占总短波比例，范围 [0, 1]")
    add_bullet(doc, "LAI_eff = min(6, LAI · clumping)")

    doc.add_heading("3.3  Noah 长波层间交换", level=2)
    add_para(doc, "冠层与地面层净辐射：")
    add_equation(doc, "R_n,canopy = SW_abs,canopy − [lw_air,c + lw_can,c · T_c⁴]", "3.3")
    add_equation(doc, "R_n,ground = SW_abs,ground − [lw_air,g + lw_can,g · T_s⁴]", "3.4")
    add_para(doc, "lw_air 与 lw_can 系数由 noah_lw_coeff_canopy / noah_lw_coeff_ground 给出，"
             "能量平衡中对 T 的导数用 noah_drn_dtemp。")

    doc.add_heading("3.4  PAR 阴阳叶分配", level=2)
    add_para(
        doc,
        "partition_par 根据两流吸收的直射/散射分量与 sunlit_frac，划分 PAR_sun、PAR_shade "
        "及 LAI_sun、LAI_shade，供双叶光合与气孔计算使用。"
    )

    doc.add_heading("3.5  离线辐射（irad = 0）", level=2)
    add_equation(doc, "R_n = SW_in (1 − α) + LW_in − ε σ T_s⁴", "3.5")
    add_para(doc, "无植被冠层时用于快速对照试验。")

    doc.add_heading("3.6  饱和水汽压与比湿", level=2)
    add_equation(doc, "e_s(T) = 611.2 exp[17.67 (T − 273.15) / ((T − 273.15) + 243.5)]  (Pa)", "3.6")
    add_equation(doc, "q_s = ε e_s / [P_a − (1 − ε) e_s]  (kg/kg)", "3.7")
    add_para(doc, "强迫读入 RH 后由 derive_humidity 计算 VPD 与 qa。")
    add_equation(doc, "VPD = max(e_s(T_a) − e_a, 0),   e_a = RH/100 · e_s(T_a)", "3.8")


def build_chapter4(doc: Document) -> None:
    doc.add_heading("第4章  边界层湍流与空气动力学阻力", level=1)
    doc.add_heading("4.1  过程简介", level=2)
    add_para(
        doc,
        "湍流模块（mod_turbulence）基于 Monin-Obukhov 相似理论，在参考高度 z_ref = 2 m "
        "计算空气动力学阻力 ra 与感热通量 H。稳定度修正采用 Dyer-Hicks 型 ψ_m、ψ_h 函数；"
        "Obukhov 长度 L 通过 3 次迭代自洽求解。"
    )

    doc.add_heading("4.2  感热通量", level=2)
    add_equation(doc, "H = ρ c_p (T_s − T_a) / r_a", "4.1")
    add_para(doc, "ρ = 1.225 kg m⁻³，c_p = 1005 J kg⁻¹ K⁻¹。")

    doc.add_heading("4.3  摩擦速度与 Obukhov 长度", level=2)
    add_equation(doc, "u_* = κ u / [ln((z_ref − d) / z_0) − ψ_m(ζ)]", "4.2")
    add_equation(doc, "L = −u_*² θ_v / (κ g θ_v*)", "4.3")
    add_equation(doc, "ζ = (z_ref − d) / L", "4.4")
    add_para(doc, "κ = 0.4 为 von Kármán 常数，d 为零平面位移，z_0 为粗糙度长度。")

    doc.add_heading("4.4  稳定度修正函数", level=2)
    add_para(doc, "不稳定条件（ζ < 0，Dyer-Hicks）：")
    add_equation(doc, "x = (1 − 16ζ)^0.25", "4.5")
    add_equation(
        doc,
        "ψ_m = 2 ln[(1+x)/2] + ln[(1+x²)/2] − 2 arctan(x) + arctan(1)",
        "4.6",
    )
    add_equation(doc, "ψ_h = 2 ln{[1 + (1 − 16ζ)^0.5] / 2}", "4.7")
    add_para(doc, "稳定条件（ζ ≥ 0）：")
    add_equation(doc, "ψ_m = ψ_h = −5ζ", "4.8")

    doc.add_heading("4.5  空气动力学阻力", level=2)
    add_equation(
        doc,
        "r_a = [ln((z_ref−d)/z_0) − ψ_m] [ln((z_ref−d)/z_0) − ψ_h] / (κ² u)",
        "4.9",
    )
    add_para(doc, "ra 下限取 1 s m⁻¹，风速下限 0.1 m s⁻¹。")


def build_chapter5(doc: Document) -> None:
    doc.add_heading("第5章  地表能量平衡", level=1)
    doc.add_heading("5.1  控制方程", level=2)
    add_equation(doc, "R_n − H − LE − G = 0", "5.1")
    add_para(
        doc,
        "当 icanopy=2 且 LAI>0.01 时，mod_surface::solve_two_leaf_noahmp 采用 Noah-MP "
        "序贯求解：先固定 Ts 迭代冠层温度 Tc（20 次内层 MO 稳定度），再固定 Tc 迭代 Ts"
        "（5 次），外层 Picard 循环 3 次。大叶模式仍用单面 Newton 迭代 Ts。"
    )

    doc.add_heading("5.2  冠层/地面通量组装", level=2)
    add_equation(doc, "H = H_canopy + H_ground", "5.2")
    add_equation(doc, "LE = LE_canopy + LE_soil", "5.3")
    add_equation(doc, "R_n = R_n,canopy + R_n,ground", "5.4")
    add_para(
        doc,
        "冠层感热/潜热用冠层空气阻力 rac、叶面边界层 rb；地面部分考虑冠层遮蔽与土壤/积雪热通量。"
        "Medlyn 气孔在冠层迭代内嵌套 photosyn_stomatal_coupling_tl，阴阳叶分别计算 rs。"
    )

    doc.add_heading("5.3  潜热与心理常数", level=2)
    add_para(doc, "潜热计算考虑饱和水汽压对温度的导数及心理常数 γ：")
    add_equation(doc, "de_s/dT = 4098 e_s / (T − 273.15 + 243.5)²", "5.5")
    add_equation(doc, "γ = c_p P_a / (ε λ)", "5.6")

    doc.add_heading("5.4  地面热通量", level=2)
    add_equation(doc, "G = G_soil + G_skin + G_snow", "5.7")
    add_para(doc, "土壤导热、地表皮肤热储量与积雪导热分别由 mod_soil_heat、skin_storage_flux、"
             "snow_ground_flux 贡献。")

    doc.add_heading("5.5  收敛判据", level=2)
    add_para(
        doc,
        "层内瞬时残差 instant_res 写入 flux%ebal_res；若 |ebal_res| > 5·tol 则打印警告。"
        "典型正午（SW=800 W/m²）残差约数十 W/m²，较旧版单面 Newton 方案显著改善。"
    )
    add_para(doc, "蒸腾速率 ET = LE / λ (m s⁻¹)。")


def build_chapter6(doc: Document) -> None:
    doc.add_heading("第6章  气孔导度", level=1)
    doc.add_heading("6.1  过程简介", level=2)
    add_para(
        doc,
        "气孔导度模块（mod_conductance）提供两种可选方案：经验 Jarvis 方案与半机理 "
        "Medlyn 方案。气孔导度 g_s 单位为 m s⁻¹（水汽），换算为冠层阻力时考虑 LAI。"
    )

    doc.add_heading("6.2  Jarvis 方案（istomatal = 1）", level=2)
    add_equation(doc, "g_s = g_s,max · f_light · f_temp · f_VPD · f_soil", "6.1")
    add_para(doc, "各胁迫因子：")
    add_equation(doc, "f_light = max[SW / (SW + 200), 0.05]", "6.2")
    add_equation(doc, "f_temp = max{min[1 − 0.0016(T_a − 298.15)², 1], 0}", "6.3")
    add_equation(doc, "f_VPD = exp(−0.001 · VPD)", "6.4")
    add_equation(doc, "f_soil = max(β, stress_hydro)", "6.5")

    doc.add_heading("6.3  Medlyn 方案（istomatal = 2）", level=2)
    add_equation(
        doc,
        "g_s = g_0 + 1.6 (1 + g_1/√VPD_kPa) · A_n / C_s",
        "6.6",
    )
    add_para(doc, "其中 A_n 为叶片净同化速率 (μmol m⁻² s⁻¹)，C_s 为叶片表面 CO₂ 浓度：")
    add_equation(doc, "C_s = CO₂ × 10⁻⁶ · P_a / (R T_a)  (mol m⁻³)", "6.7")

    doc.add_heading("6.4  冠层阻力", level=2)
    add_equation(doc, "r_s = 1 / (LAI · g_s)    (LAI > 0.01)", "6.8")
    add_para(doc, "g_s 下限 10⁻⁶ m s⁻¹。")

    doc.add_heading("6.5  An–gs–Ci 耦合迭代", level=2)
    add_para(
        doc,
        "Medlyn 模式下，能量平衡内调用 photosyn_stomatal_coupling，进行 3 次固定点迭代："
        "(1) 由 Ci 计算 Farquhar 光合得 An；(2) 由 An 计算 gs；(3) 由 gs 更新 Ci："
    )
    add_equation(doc, "C_i = C_a − A_n · 1.6 / g_s · P_a / (R T_a)", "6.9")
    add_para(doc, "Ci 下限 10 μmol mol⁻¹。")


def build_chapter7(doc: Document) -> None:
    doc.add_heading("第7章  光合作用与呼吸", level=1)
    doc.add_heading("7.1  Farquhar C3 模型", level=2)
    add_para(
        doc,
        "光合模块（mod_photosyn）实现 Farquhar et al. C3 光合作用模型，Rubisco 限制与"
        "电子传递限制取最小值，再减去暗呼吸得到净同化。"
    )
    add_equation(
        doc,
        "W_c = V_cmax (C_i − Γ*) / [C_i + K_c (1 + O/K_o)]",
        "7.1",
    )
    add_equation(
        doc,
        "W_j = J (C_i − Γ*) / (C_i + 2Γ*)",
        "7.2",
    )
    add_equation(doc, "A_n = min(W_c, W_j) − R_d", "7.3")
    add_equation(doc, "GPP = A_n · LAI", "7.4")
    add_para(doc, "模型常数：K_c = 404.9，K_o = 278400，Γ* = 42.75 μmol mol⁻¹，O = 210000 μmol mol⁻¹。")

    doc.add_heading("7.2  电子传递速率与 PAR", level=2)
    add_para(doc, "双叶模式（icanopy=2）下 PAR 来自 mod_radtran 的 PAR_sun / PAR_shade；"
             "大叶模式仍用 PAR_abs ≈ 0.5·SW 近似。")
    add_equation(doc, "J = J_max · PAR_abs / (PAR_abs + 150)", "7.5")
    add_equation(doc, "GPP = GPP_sun + GPP_shade    (双叶模式)", "7.6")

    doc.add_heading("7.3  温度响应", level=2)
    add_para(doc, "Vcmax 采用 Bernacchi 型活化/失活（mod_acclimation::vcmax_at_temp）：")
    add_equation(
        doc,
        "V_cmax(T) = V_cmax,25 · exp[Ha(T−298.15)/(T·298.15·R)] / {1 + exp[(Sv·T−Hd)/(R·T)]}",
        "7.7",
    )
    add_para(doc, "Ha = 65330 J mol⁻¹，Hd = 149250 J mol⁻¹，Sv = 485 J mol⁻¹ K⁻¹。")
    add_equation(
        doc,
        "J_max(T) = J_max,25 · exp[37000(T−298.15)/(T·298.15·R)]",
        "7.8",
    )
    add_equation(doc, "R_d(T) = R_d,25 · 2^((T−298.15)/10)", "7.9")

    doc.add_heading("7.4  水分与养分胁迫", level=2)
    add_equation(doc, "V_cmax, J_max ← V_cmax, J_max · f_stress · ν_stress · π_stress", "7.10")
    add_para(doc, "f_stress 来自植物水力模块；ν_stress、π_stress 来自 CNP 模块。")

    doc.add_heading("7.5  EEO 光合适应", level=2)
    add_para(
        doc,
        "当 eeo_on = .true. 时，mod_acclimation 对气温做指数平滑（时间尺度 accl_days），"
        "调节基准 Vcmax25 与 Rd25（Ren et al. 2025 简化实现）："
    )
    add_equation(doc, "T_accl ← (1−w) T_accl + w T_a,   w = Δt / (τ + Δt)", "7.11")
    add_equation(doc, "V_cmax,25 ← V_cmax · exp[0.04(T_accl − 298.15)],  限幅 [0.5, 2.0]", "7.12")
    add_equation(doc, "R_d,25 ← R_d,base · exp[0.06(T_accl − 298.15)],  限幅 [0.4, 2.5]", "7.13")


def build_chapter8(doc: Document) -> None:
    doc.add_heading("第8章  植物水力学（PHS）", level=1)
    doc.add_heading("8.1  过程简介", level=2)
    add_para(
        doc,
        "植物水力模块（mod_planthydro）实现简化 Kennedy et al. (2019) 风格植物水力系统（PHS）。"
        "当 ihydro = 2 时，蒸腾水分胁迫由水力胁迫 stress_hydro 与土壤水分因子 β 共同决定；"
        "Richards 土壤水方案中根系吸水与 psi_root 耦合。"
    )

    doc.add_heading("8.2  木质部导度与脆弱性曲线", level=2)
    add_equation(
        doc,
        "V(P) = 1 / {1 + [|ψ − P50| / |P50|]^Ck}    (ψ < P50)",
        "8.1",
    )
    add_equation(doc, "k_xylem = k_x,max · max(V, 0.01)", "8.2")

    doc.add_heading("8.3  水势动力学", level=2)
    add_para(doc, "有蒸腾时，沿土壤–根–木–叶链水势下降：")
    add_equation(doc, "ET_m = ET / ρ_w", "8.3")
    add_equation(doc, "Δψ = min(ET_m / k_eff · 10⁻⁸, 0.8 MPa)", "8.4")
    add_equation(doc, "ψ_leaf = max(ψ_soil − Δψ, −3 MPa)", "8.5")
    add_equation(doc, "ψ_xylem = max(ψ_soil − 0.6Δψ, −3.5 MPa)", "8.6")
    add_equation(doc, "ψ_root = max(ψ_soil − 0.2Δψ, −3 MPa)", "8.7")
    add_para(doc, "无蒸腾时水势向 ψ_soil 缓慢恢复。")
    add_para(doc, "土壤水势为根区加权平均：")
    add_equation(doc, "ψ_soil = Σ_i f_root,i · ψ_soil,i", "8.8")

    doc.add_heading("8.4  水力胁迫因子", level=2)
    add_equation(
        doc,
        "stress_hydro = 1                          (ψ_leaf ≥ ψ50,leaf)",
        "8.9",
    )
    add_equation(
        doc,
        "stress_hydro = max[0, (ψ_leaf − 2ψ50)/(ψ50 − 2ψ50)]    (ψ_leaf < ψ50,leaf)",
        "8.10",
    )

    doc.add_heading("8.5  蒸腾水分胁迫综合因子", level=2)
    add_equation(doc, "f_stress = max(stress_hydro · β, 0.3β)    (PHS 模式)", "8.11")
    add_equation(doc, "f_stress = β    (beta 模式)", "8.12")
    add_para(doc, "f_stress 限幅 [0.1, 1.0]。")


def build_chapter9(doc: Document) -> None:
    doc.add_heading("第9章  土壤温度", level=1)
    doc.add_heading("9.1  土壤垂向离散", level=2)
    add_para(
        doc,
        "土壤热模块（mod_soil_heat）将 0–2 m 土壤柱分为 nsoil 层（默认 6 层），"
        "层厚指数递增后归一化至总深 2 m。层中心深度 z_mid 由层厚累积计算。"
    )

    doc.add_heading("9.2  地面热通量", level=2)
    add_equation(doc, "G = k_s (T_soil,1 − T_s) / (Δz_1/2)", "9.1")
    add_para(doc, "|G| ≤ 300 W/m²。")

    doc.add_heading("9.3  土壤温度更新", level=2)
    add_para(doc, "顶层采用显式能量收支：")
    add_equation(doc, "T_soil,1 ← T_soil,1 − G · Δt / (C_s · Δz_1)", "9.2")
    add_para(doc, "C_s 为土壤体积热容量 (J m⁻³ K⁻¹)。深层采用松弛扩散：")
    add_equation(doc, "T_soil,i ← T_soil,i + 0.05 (T_soil,i−1 − T_soil,i),   i ≥ 2", "9.3")
    add_para(
        doc,
        "当前版本为 MVP 实现：未包含冻土相变、积雪隔热与完整热传导方程数值求解。"
        "青藏高原冻土–植被耦合研究需在此基础上扩展相变潜热项。"
    )


def build_chapter10(doc: Document) -> None:
    doc.add_heading("第10章  土壤水分", level=1)
    doc.add_heading("10.1  过程简介", level=2)
    add_para(
        doc,
        "土壤水模块（mod_soil_water）提供单库 Bucket 与多层 Richards 两种方案。"
        "Richards 方案采用 van Genuchten 土壤水力特性曲线，考虑入渗、层间达西通量、"
        "根系吸水与底层自由排水。"
    )

    doc.add_heading("10.2  土壤水分胁迫因子 β", level=2)
    add_para(doc, "Richards 模式：")
    add_equation(doc, "θ_avg = Σ θ_i Δz_i / Σ Δz_i", "10.1")
    add_equation(doc, "β = (θ_avg − θ_r) / (θ_s − θ_r),   限幅 [0, 1]", "10.2")
    add_para(doc, "Bucket 模式：W 在 W_wilt 与 W_field 之间线性映射。")

    doc.add_heading("10.3  Bucket 方案", level=2)
    add_equation(doc, "ΔW = P − ET · Δt / 3600    (mm)", "10.3")
    add_equation(doc, "W ← min(W + ΔW, W_field)", "10.4")

    doc.add_heading("10.4  van Genuchten 特征曲线", level=2)
    add_equation(doc, "m = 1 − 1/n", "10.5")
    add_equation(doc, "S_e = [1 + (α|ψ|)^n]^(−m)    (ψ < 0)", "10.6")
    add_equation(doc, "θ = θ_r + (θ_s − θ_r) S_e", "10.7")
    add_equation(doc, "ψ(θ) = −[(S_e^(−1/m) − 1)^(1/n)] / α", "10.8")

    doc.add_heading("10.5  非饱和导水率", level=2)
    add_equation(
        doc,
        "K(θ) = K_sat S_e^0.5 [1 − (1 − S_e^(1/m))^m]²",
        "10.9",
    )

    doc.add_heading("10.6  Richards 方程离散", level=2)
    add_para(doc, "层间达西通量（含重力项）：")
    add_equation(doc, "q_{i+1} = K_mid [(ψ_{i+1} − ψ_i)/Δz_i + 1]", "10.10")
    add_para(doc, "含水率更新：")
    add_equation(doc, "θ_i ← θ_i + (q_i − q_{i+1} − sink_i) Δt / Δz_i", "10.11")

    doc.add_heading("10.7  根系分布与吸水", level=2)
    add_equation(doc, "f_root,i ∝ exp(−2 z_i),   归一化", "10.12")
    add_equation(
        doc,
        "sink_i = f_root,i · max[k_r (ψ_soil,i − ψ_root), 0] · 10⁶",
        "10.13",
    )
    add_para(doc, "各层 sink 按比例分配总蒸腾 ET。")


def build_chapter11(doc: Document) -> None:
    doc.add_heading("第11章  土壤碳循环", level=1)
    doc.add_heading("11.1  过程简介", level=2)
    add_para(
        doc,
        "土壤碳模块（mod_soil_carbon）提供 Q10 经验方案与显式微生物库方案。"
        "碳库包括凋落物 Clit、微生物 Cmic 与土壤有机质 Csom。"
    )

    doc.add_heading("11.2  温度响应", level=2)
    add_equation(doc, "f_T = 2^((T_soil − 288.15)/10)", "11.1")

    doc.add_heading("11.3  Q10 方案（isoilcarbon = 1）", level=2)
    add_equation(doc, "R_soil = 2.0 · f_T · β", "11.2")

    doc.add_heading("11.4  微生物方案（isoilcarbon = 2）", level=2)
    add_equation(doc, "decomp = k_litter · C_lit · f_T · β", "11.3")
    add_equation(doc, "R_mic = 0.1 · C_mic · f_T", "11.4")
    add_equation(doc, "R_som = k_som · C_som · f_T · 0.1", "11.5")
    add_equation(doc, "R_soil = decomp + R_mic + R_som", "11.6")

    doc.add_heading("11.5  碳库更新", level=2)
    add_equation(doc, "C_lit ← C_lit − decomp·Δt/86400 + litterfall·Δt/86400", "11.7")
    add_equation(doc, "growth = CUE · decomp", "11.8")
    add_equation(doc, "C_mic ← C_mic + (growth − 0.1·C_mic·f_T)·Δt/86400", "11.9")
    add_equation(doc, "C_som ← C_som + (1−CUE)·decomp·0.5·Δt/86400", "11.10")

    doc.add_heading("11.6  生态系统碳通量诊断", level=2)
    add_equation(doc, "R_eco = R_leaf + R_soil", "11.11")
    add_equation(doc, "NEE = R_eco − GPP", "11.12")
    add_para(doc, "NEE > 0 表示碳源，NEE < 0 表示碳汇。")


def build_chapter12(doc: Document) -> None:
    doc.add_heading("第12章  碳氮磷（CNP）耦合", level=1)
    doc.add_heading("12.1  过程简介", level=2)
    add_para(
        doc,
        "CNP 模块（mod_cnp）实现简化 CASA 风格养分限制。icnp 开关控制碳（C）、"
        "碳氮（CN）、碳氮磷（CNP）三个层级。"
    )
    add_table(
        doc,
        ["icnp", "行为"],
        [
            ["0", "关闭养分循环，ν_stress = π_stress = 1"],
            ["1", "仅碳，无养分限制"],
            ["2", "CN：GPP 消耗 N_lab，ν_stress 反馈光合"],
            ["3", "CNP：额外磷库 P_lab，π_stress 反馈光合"],
        ],
    )

    doc.add_heading("12.2  氮循环", level=2)
    add_equation(doc, "N_demand = GPP · 10⁻⁶ · Δt / 86400", "12.1")
    add_equation(doc, "N_uptake = min(N_uptake,max · Δt/86400 · β, 0.1 · N_lab)", "12.2")
    add_equation(doc, "N_lab ← max(N_lab − N_demand + N_uptake, 0.1)", "12.3")
    add_equation(doc, "ν_stress = min(1, N_lab / 5.0)", "12.4")

    doc.add_heading("12.3  磷循环", level=2)
    add_equation(doc, "P_demand = GPP · 10⁻⁷ · Δt / 86400", "12.5")
    add_equation(doc, "P_uptake = min(P_uptake,max · Δt/86400 · β, 0.1 · P_lab)", "12.6")
    add_equation(doc, "P_lab ← max(P_lab − P_demand + P_uptake, 0.05)", "12.7")
    add_equation(doc, "π_stress = min(1, P_lab / 0.5)", "12.8")


def build_chapter13(doc: Document) -> None:
    doc.add_heading("第13章  守恒检查与诊断", level=1)
    doc.add_heading("13.1  能量平衡残差", level=2)
    add_equation(doc, "resid_E = R_n − H − LE − G", "13.1")
    add_para(doc, "当 |resid_E| > 5 W/m² 且 check_conservation = .true. 时，每 48 步打印警告。")

    doc.add_heading("13.2  水分平衡残差", level=2)
    add_equation(doc, "resid_W = P − ET · Δt / 3600    (mm)", "13.2")

    doc.add_heading("13.3  输出变量", level=2)
    add_table(
        doc,
        ["变量", "含义", "单位"],
        [
            ["Rn, H, LE, G", "辐射、感热、潜热、地面热通量", "W/m²"],
            ["Ts", "地表温度", "K"],
            ["beta", "土壤水分胁迫", "0–1"],
            ["W", "根区/剖面水量", "mm"],
            ["GPP, NEE", "总初级生产、净生态系统交换", "μmol/m²/s"],
            ["psi_leaf", "叶水势", "MPa"],
            ["stress_hydro", "水力胁迫因子", "0–1"],
        ],
    )


def build_appendices(doc: Document) -> None:
    doc.add_heading("附录 A  物理常数", level=1)
    add_table(
        doc,
        ["符号", "数值", "单位", "说明"],
        [
            ["c_p", "1005", "J kg⁻¹ K⁻¹", "空气定压比热"],
            ["ρ", "1.225", "kg m⁻³", "空气密度（海平面参考）"],
            ["σ", "5.67×10⁻⁸", "W m⁻² K⁻⁴", "Stefan-Boltzmann 常数"],
            ["λ", "2.45×10⁶", "J kg⁻¹", "汽化潜热"],
            ["κ", "0.4", "—", "von Kármán 常数"],
            ["g", "9.81", "m s⁻²", "重力加速度"],
            ["R_d", "287", "J kg⁻¹ K⁻¹", "干空气气体常数"],
            ["ε", "0.622", "—", "水汽/干空气分子量比"],
            ["T_frz", "273.15", "K", "冰点"],
        ],
    )

    doc.add_heading("附录 B  Namelist 主要参数", level=1)
    doc.add_heading("B.1  config_nml", level=2)
    add_table(
        doc,
        ["变量", "默认值", "说明"],
        [
            ["dt", "1800.0", "时间步长 (s)"],
            ["nspinup", "1", "spin-up 循环次数"],
            ["nsoil", "6", "土壤层数"],
            ["carbon_on", ".false.", "碳循环开关"],
            ["forcing_file", "data/nc/sample_forcing.nc", "强迫路径（txt 或 nc）"],
            ["soil_init_file", "data/nc/sample_soil_init.nc", "土壤初值路径"],
            ["max_iter", "50", "能量平衡最大迭代（大叶模式）"],
            ["tol", "5.0", "能量平衡收敛阈值 (W/m²)"],
            ["icanopy", "2", "冠层方案：1=大叶，2=双叶"],
            ["irad", "1", "辐射：0=离线，1=两流，2=GCM"],
            ["isnow", "1", "积雪开关"],
            ["ifrost", "1", "冻土相变开关"],
            ["istomatal", "2", "气孔方案"],
            ["ihydro", "2", "水分胁迫方案"],
            ["isoilwater", "2", "土壤水方案"],
            ["isoilcarbon", "2", "土壤碳方案"],
            ["icnp", "3", "CNP 方案"],
            ["eeo_on", ".true.", "EEO 光合适应"],
            ["accl_days", "30.0", "适应时间尺度 (day)"],
        ],
    )

    doc.add_heading("B.2  param_nml 主要分组", level=2)
    add_bullet(doc, "地表/植被：z0, zdisp, albedo, emiss, hc, lai")
    add_bullet(doc, "气孔：gs_max（Jarvis），g0, g1（Medlyn）")
    add_bullet(doc, "光合：vcmax, jmax, rd25_base")
    add_bullet(doc, "土壤水力：ksat, poros, theta_r, alpha_vg, n_vg, W_field, W_wilt")
    add_bullet(doc, "土壤热：soil_heat_cap, soil_cond")
    add_bullet(doc, "植物水力：p50_xylem, ck_xylem, kx_max, kr_max, psi50_leaf")
    add_bullet(doc, "土壤碳：cue_micro, k_litter, k_som, litterfall")
    add_bullet(doc, "养分：n_uptake_max, p_uptake_max")

    doc.add_heading("附录 C  强迫与输出文件格式", level=1)
    doc.add_heading("C.1  文本强迫", level=2)
    add_para(doc, "首行：# SW LW Ta P WS PA CO2 RH [cos_sza sw_beam_frac]")
    add_para(doc, "8 列格式向后兼容；后两列缺省时由 infer_sw_geometry 从 SW 推断。")
    doc.add_heading("C.2  NetCDF 强迫", level=2)
    add_para(doc, "维度 time；变量 SW, LW, Ta, P, WS, PA, CO2, RH；可选 cos_sza, sw_beam_frac。")
    add_para(doc, "生成：make forcing-nc 或 python3 scripts/txt_to_nc_forcing.py")
    doc.add_heading("C.3  NetCDF 输出", level=2)
    add_para(doc, "默认 results/nc/output.nc，23 个诊断量含 GPP_sun, GPP_shade, albedo_eff, snow_swe 等。")
    add_para(doc, "缺测值标志：-9999（读入后替换为合理默认值）。")

    doc.add_heading("附录 D  参考文献", level=1)
    refs = [
        "Farquhar G.D., von Caemmerer S., Berry J.A. (1980). A biochemical model of photosynthetic CO₂ assimilation in leaves of C3 species. Planta, 149, 78–90.",
        "Medlyn B.E. et al. (2011). Reconciling the optimal and empirical approaches to modelling stomatal conductance. Global Change Biology, 17, 2134–2144.",
        "van Genuchten M.T. (1980). A closed-form equation for predicting the hydraulic conductivity of unsaturated soils. Soil Science Society of America Journal, 44, 892–898.",
        "Kennedy D. et al. (2019). Implementing plant hydraulics in the Community Land Model, version 5. Journal of Advances in Modeling Earth Systems, 11, 485–513.",
        "Ren C. et al. (2025). EEO photosynthetic acclimation (simplified implementation in mod_acclimation).",
        "Niu G.-Y. et al. (2011). The community Noah land surface model with multiparameterization options (Noah-MP). Journal of Geophysical Research, 116, D12109.",
        "Bonan G.B. (2019). Climate Change and Terrestrial Ecosystem Modeling. Cambridge University Press. (CLM 技术文档结构参考)",
    ]
    for ref in refs:
        add_bullet(doc, ref)

    doc.add_heading("附录 E  已知限制", level=1)
    limits = [
        "单点 0-D：无完整多层冠层（icanopy=3 未实现）、GCM 耦合为占位接口。",
        "积雪、冻土、PHS、Richards、微生物碳、CNP 均为 MVP 级，需青藏高原站点标定。",
        "能量平衡残差在部分时步仍可达 40–250 W/m²，与序贯求解及松耦合有关。",
        "碳循环 spin-up 仅循环短期强迫，未实现多年气候序列碳库平衡。",
        "叶光学参数（rho_leaf, tau_leaf）目前硬编码于 mod_radtran，未暴露 namelist。",
    ]
    for lim in limits:
        add_bullet(doc, lim)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    set_doc_styles(doc)
    add_page_number_footer(doc)

    build_title_page(doc)

    doc.add_heading("目录", level=1)
    add_toc(doc)
    add_page_break(doc)

    build_chapter1(doc)
    add_page_break(doc)
    build_chapter2(doc)
    add_page_break(doc)
    build_chapter3(doc)
    add_page_break(doc)
    build_chapter4(doc)
    add_page_break(doc)
    build_chapter5(doc)
    add_page_break(doc)
    build_chapter6(doc)
    add_page_break(doc)
    build_chapter7(doc)
    add_page_break(doc)
    build_chapter8(doc)
    add_page_break(doc)
    build_chapter9(doc)
    add_page_break(doc)
    build_chapter10(doc)
    add_page_break(doc)
    build_chapter11(doc)
    add_page_break(doc)
    build_chapter12(doc)
    add_page_break(doc)
    build_chapter13(doc)
    add_page_break(doc)
    build_appendices(doc)

    return doc


def main() -> None:
    doc = build_document()
    doc.save(str(OUTPUT))
    print(f"Technical manual written to: {OUTPUT}")
    print(f"Pages/paragraphs: {len(doc.paragraphs)} paragraphs")
    print("Open in Word/LibreOffice and right-click TOC -> Update Field for page numbers.")


if __name__ == "__main__":
    main()